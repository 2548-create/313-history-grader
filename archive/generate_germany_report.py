#!/usr/bin/env python3
"""生成德国分裂原因批改报告"""

import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_colored_paragraph(doc, text, red_keywords=None, score_note=None):
    p = doc.add_paragraph()
    if not red_keywords:
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        segments = [text]
        for error_word, correct_word in red_keywords:
            new_segments = []
            for seg in segments:
                if isinstance(seg, tuple):
                    new_segments.append(seg)
                    continue
                parts = seg.split(error_word)
                for i, part in enumerate(parts):
                    if part:
                        new_segments.append(part)
                    if i < len(parts) - 1:
                        new_segments.append((error_word, correct_word))
            segments = new_segments
        for seg in segments:
            if isinstance(seg, tuple):
                error_word, correct_word = seg
                run = p.add_run(error_word)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
                if correct_word:
                    note = p.add_run(f"（应为{correct_word}）")
                    note.font.color.rgb = RGBColor(255, 0, 0)
                    note.font.size = Pt(9)
            else:
                run = p.add_run(seg)
                run.font.color.rgb = RGBColor(0, 0, 0)
    if score_note:
        note_run = p.add_run(f"  ({score_note})")
        note_run.font.color.rgb = RGBColor(128, 128, 128)
        note_run.font.size = Pt(9)
    return p

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10)

    # 标题
    title = doc.add_heading('313历史学统考论述题批改报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 题目
    doc.add_heading('【题目】', level=1)
    doc.add_paragraph('论述中世纪德国未能形成统一国家的原因')

    # 总分概览
    doc.add_heading('【总分概览】', level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    headers = ['维度', '得分', '满分']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    data = [
        ['要点踩点', '18', '30'],
        ['论述组织（含卷面）', '8.5', '10'],
        ['合计', '26.5/40', '']
    ]
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.rows[i + 1].cells[j].text = cell
    doc.add_paragraph()

    # 学生答案
    doc.add_heading('【学生答案】', level=1)
    student_text = """24. 中世纪德国长时间处于分裂状态，是由皇位继承、内部争斗、教皇阻挠等多方面因素综合形成的结果，并对欧洲政治格局产生深远影响。原因如下。
一、德国独特的皇位继承制度使皇权式微。
中世纪德国的皇帝由选帝侯选举产生且选帝侯为维护自身权益常选举实力较弱的诸侯当选为德国皇帝，德皇没有足够权威进行国家统一是中世纪德国长时期分裂的重要因素。
二、《黄金诏书》使德国分裂法律化，延缓德国统一。
黄金诏书是神圣罗马帝国皇帝查理四世颁布，从法律上确立了德意志侯国的分裂体制，进一步削弱皇权，加剧德意志的政治分裂，阻碍德意志迈向民族国家的进程，迟滞德国的统一。
三、各诸侯实力均等且长期混战，无绝对实力足以统一的诸侯。
德意志各诸侯国长期混战，因宗教、他国干涉等因素，彼此混战消耗国力，加剧矛盾且未有超级强国的选帝侯完成统一，各国实力均衡，故相互林立，制衡，形成长期对峙局面，未能统一。
四、教皇的干涉与世俗权力之争，消耗德国国力。
教皇长期干涉德意志的统一，因德意志分裂有利于教皇控制德国，且德国长期受教皇剥削严重，被称为"教皇的奶牛"，故教皇对德国统一加以阻挠，且奥托一世卡诺莎之辱，教皇长期与德国争斗，消耗德国国力。
五、德皇重心转移于意大利，使德国统一受影响。
德国自神圣罗马帝国，历任德皇将重心置于意大利，沉溺于罗马加冕，故历任德皇时期消耗国力且对德国内部关注度不足，是德国分裂原因。
六、欧洲大国的干涉，阻碍德意志统一
法国的大陆称霸政策，法国将德国看作欧洲大陆最大称霸阻碍，故法国外交分化政策，使德意志诸侯维持平衡，难以一统。
综上所述，中世纪德国未能形成统一国家是内外双重因素结合的结果，使德国长期分裂，对欧洲的格局具有重大影响。"""

    red_kw = [("奥托一世", "亨利四世")]
    for para in student_text.split('\n'):
        para = para.strip()
        if para:
            score_note = None
            if '(得分:' in para:
                import re
                match = re.search(r'\(得分:[^)]+\)', para)
                if match:
                    score_note = match.group()[1:-1]
                    para = para[:match.start()] + para[match.end():]
            add_colored_paragraph(doc, para, red_keywords=red_kw, score_note=score_note)

    note_p = doc.add_paragraph()
    note_run = note_p.add_run("注：标红处为史实错误，括号内为正确表述。")
    note_run.font.color.rgb = RGBColor(128, 128, 128)
    note_run.font.size = Pt(9)

    # 踩分点对照表
    doc.add_heading('【踩分点对照表】', level=1)

    sections_data = [
        {
            "name": "政治制度", "max": 12, "score": 6,
            "points": [
                ["皇位选举制导致皇权先天不足", "皇帝由选帝侯选举产生；选帝侯为维护自身权益常选举实力较弱的诸侯当选皇帝", "4/4", "完整答出选举制原理及选帝侯动机，给满分"],
                ["大空位时代与诸侯割据固化", "无", "0/4", "未提及1250年霍亨斯陶芬王朝绝嗣、大空位时代（1254-1273年）这一关键史实，完全缺失"],
                ["封建领地分裂格局", "各诸侯实力均等且长期混战，相互林立制衡", "2/2", "答到诸侯割据和实力均等，给满分"],
                ["缺乏统一的经济基础与民族认同", "无", "0/2", "完全缺失经济和文化认同层面的分析"]
            ]
        },
        {
            "name": "法律固化", "max": 6, "score": 5,
            "points": [
                ["金玺诏书的颁布及其内容", "黄金诏书是查理四世颁布；从法律上确立了德意志侯国的分裂体制，削弱皇权", "3/4", "基本命中，缺少1356年及诸侯具体特权内容，属部分命中扣1分"],
                ["金玺诏书的历史影响", "使德国分裂法律化，延缓德国统一，阻碍迈向民族国家", "2/2", "答到法律固化和延缓统一，给满分"]
            ]
        },
        {
            "name": "宗教势力", "max": 6, "score": 4,
            "points": [
                ["叙任权斗争", "奥托一世卡诺莎之辱", "1/3", "史实硬伤：卡诺莎事件应为亨利四世（1077年），而非奥托一世。只答到名称未展开，属触及边缘扣2分"],
                ["德皇南倾意大利", "德皇重心转移于意大利，历任德皇将重心置于意大利，沉溺于罗马加冕，消耗国力", "3/3", "完整答出德皇南倾意大利的史实和影响，给满分"]
            ]
        },
        {
            "name": "外部干涉", "max": 6, "score": 3,
            "points": [
                ["法国的分化政策", "法国的大陆称霸政策，法国外交分化政策使德意志诸侯维持平衡", "3/4", "答出法国分化政策但缺少具体策略和蚕食领土等细节，属基本命中扣1分"],
                ["其他外部势力的影响", "因宗教、他国干涉等因素", "0/2", "仅一笔带过未具体提及任何外部势力，属触及边缘扣2分"]
            ]
        }
    ]

    for sec in sections_data:
        p = doc.add_paragraph(f"（{sec['name']}）（满分{sec['max']}分，得分{sec['score']}分）")
        p.runs[0].bold = True

        tbl = doc.add_table(rows=len(sec['points'])+1, cols=4)
        tbl.style = 'Light Grid Accent 1'
        for i, h in enumerate(['踩分点', '学生作答', '得分/满分', '点评']):
            tbl.rows[0].cells[i].text = h
            tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for i, row in enumerate(sec['points']):
            for j, txt in enumerate(row):
                tbl.rows[i+1].cells[j].text = txt
                for para in tbl.rows[i+1].cells[j].paragraphs:
                    for run in para.runs:
                        run.bold = False
        doc.add_paragraph()

    # 论述组织评分
    doc.add_heading('【论述组织评分】（满分10分）', level=1)
    lang_data = [
        ['整体结构', '2/2', '结构完整，有总起分述总结，六个板块层次清晰'],
        ['论证逻辑', '2/2', '论证逻辑基本连贯，各点之间有递进关系'],
        ['史论结合', '1/2', '有史实但部分关键史实缺失，史论结合一般'],
        ['文字表达', '1/1', '表述基本通顺，有1处明显史实错误'],
        ['史观运用', '0.5/1', '有历史视角但分析偏于表面，缺乏对制度性原因的深入揭示'],
        ['卷面整洁', '2/2', '字迹工整清晰，卷面整洁'],
        ['合计', '8.5/10', '']
    ]
    tbl = doc.add_table(rows=len(lang_data), cols=3)
    tbl.style = 'Light Grid Accent 1'
    for i, row in enumerate(lang_data):
        for j, txt in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.text = txt
            if i == 0 or i == len(lang_data)-1:
                cell.paragraphs[0].runs[0].bold = True
            else:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = False
    doc.add_paragraph()

    # 本题诊断
    doc.add_heading('【本题诊断】', level=1)
    doc.add_paragraph('一、亮点', style='Heading 2')
    for item in ['结构意识强：采用总分总结构，开篇有概述，结尾有总结，六个板块层次分明，符合论述题规范',
                 '金玺诏书要点基本命中：答出查理四世颁布、分裂法律化、延缓统一等核心内容']:
        p = doc.add_paragraph()
        p.add_run(f"• {item}")

    doc.add_paragraph('二、不足', style='Heading 2')
    for item in ['史实硬伤：卡诺莎事件人物张冠李戴——应为亨利四世（1077年）而非奥托一世',
                 '知识遗漏严重——大空位时代（1254-1273年）完全缺失：这是德意志分裂制度化的关键转折',
                 '知识遗漏——经济基础和民族认同层面完全缺失',
                 '叙任权斗争细节不足：未提及格里高利七世、沃尔姆斯宗教协定（1122年）等关键史实',
                 '外部干涉分析薄弱：仅泛泛提及其他国干涉，未具体列举波兰、匈牙利等']:
        p = doc.add_paragraph()
        p.add_run(f"• {item}")

    doc.add_paragraph('三、改进方向', style='Heading 2')
    for item in ['准确记忆关键人物与事件——卡诺莎事件的主角是亨利四世（1077年），大空位时代（1254-1273年）是德意志分裂的关键转折点',
                 '补全分析维度——原因类题目应从政治制度、法律固化、宗教势力、经济基础、文化认同、外部环境等多个维度展开',
                 '加强关键史实记忆——叙任权斗争涉及格里高利七世、亨利四世、沃尔姆斯宗教协定（1122年）等核心知识点',
                 '注重史论结合深度——不仅描述现象，更要分析制度性原因']:
        p = doc.add_paragraph()
        p.add_run(f"• {item}")
    doc.add_paragraph()

    # 整体点评
    doc.add_heading('【整体点评】', level=1)
    diagnosis = "本题得分为26.5/40，内容不完整，有多处史实错误，分析偏于表面，属三类卷。学生掌握了皇位选举制、金玺诏书、德皇南倾意大利、法国分化政策等基本知识点，结构较为完整；主要失分在于大空位时代完全缺失、经济文化维度空白、叙任权斗争细节不足且存在史实硬伤（奥托一世应为亨利四世）。提分空间较大，若能在史实记忆和分析维度上加以补足，有望达到二类卷水平。"
    doc.add_paragraph(diagnosis)

    # 参考答案
    doc.add_heading('【参考答案】', level=1)
    ref_text = """中世纪德意志长期处于分裂状态，未能形成统一的民族国家，这是欧洲政治史上的重要现象。神圣罗马帝国名义上覆盖中欧大地，实则由数百个诸侯国、自由城市和教会领地拼合而成。这一分裂局面的形成，根植于其独特的政治制度设计，又经法律文件固化，并被宗教势力与外部强国不断加剧。

一、政治制度层面（12分）

（一）皇位选举制导致皇权先天不足（4分）
神圣罗马帝国的皇帝并非血缘世袭，而是由选帝侯选举产生。这一制度从建国之初便埋下了皇权脆弱的种子——每位新皇帝登基前都必须向选帝侯做出大量让步，以换取其支持。选举制使皇权始终建立在诸侯的恩赐之上，而非基于神授或血缘的绝对合法性。每次皇位更替都是一次权力再分配，诸侯借此不断蚕食中央权力。

（二）大空位时代与诸侯割据的固化（4分）
1250年霍亨斯陶芬王朝绝嗣后，德意志陷入"大空位时代"（1254-1273年），长达二十余年的无政府状态使诸侯彻底摆脱了皇帝的控制。此后的皇位争夺演变为诸侯间的混战，没有哪个势力足够强大以重新统一德意志。诸侯割据的局面因长期混战而固化，任何一个诸侯都没有实力完成统一大业。

（三）封建领地分裂格局（2分）
德意志的封建制度与法国、英国有所不同。法兰克王国拆分后，德意志地区形成了众多半独立的公国、侯国、主教区和自由城市。这些政治实体各自为政，拥有独立的司法、税收和军事权力。

（四）缺乏统一的经济基础与民族认同（2分）
德意志地区商业和城市的发展晚于西欧其他国家。德意志各地区方言差异巨大，缺乏统一的文化认同，难以形成支撑统一国家的民族意识。

二、法律层面（6分）

（一）《金玺诏书》的颁布及其内容（4分）
1356年，查理四世颁布《金玺诏书》，这是一部从根本上固化德意志分裂格局的法律文件。诏书正式确认七大选帝侯选举皇帝的制度，并赋予诸侯在其领地内近乎完全的自治权：行政权、司法权、铸币权、采矿权等，皇帝无权干涉。

（二）《金玺诏书》的历史影响（2分）
这部诏书被后世称为"德意志的小宪法"，它将分裂合法化、永久化。从1356年到1806年神圣罗马帝国灭亡，四百五十年间德意志的政治格局基本未变。

三、宗教层面（6分）

（一）叙任权斗争消耗皇权（3分）
11世纪的叙任权斗争是德意志分裂的重要催化剂。教皇格里高利七世与亨利四世的对抗中，亨利四世于1077年前往卡诺莎请求教皇赦免，这一事件极大地损害了皇权的神圣性。1122年《沃尔姆斯宗教协定》达成，皇帝失去了对主教的叙任权，教会势力在德国大幅扩张。

（二）德皇南倾意大利的战略失误（3分）
自奥托大帝以来，历代德皇都将精力和资源倾注于意大利事务，试图控制富庶的北意大利城邦。这种"南倾"战略使德意志本土长期被忽视，皇帝的注意力不在国内统一，而在跨越阿尔卑斯山的远征。

四、外部环境层面（6分）

（一）法国的分化政策（4分）
法国自卡佩王朝以来便奉行"分而治之"策略，将统一的德意志视为法国扩张的最大障碍。法国通过外交手段、金钱收买和军事干预，不断支持德意志内部的诸侯对抗皇帝，还逐步蚕食莱茵兰等德意志西部地区。

（二）其他外部势力的影响（2分）
波兰、匈牙利、波希米亚等周边势力也与德意志诸侯存在复杂的利益关系，时常介入德意志内部事务。威尼斯、热那亚等意大利城邦也与德皇存在竞争。

总之，中世纪德国未能形成统一国家，是制度缺陷、法律固化、宗教干涉和外部压力共同作用的结果。皇位选举制使皇权先天不足，《金玺诏书》将分裂合法化，教皇与德皇的长期斗争消耗了中央权威，而法国等大国的分化政策则使统一始终难以实现。这一分裂格局延续至近代，深刻影响了欧洲的政治版图。"""

    for line in ref_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('一、') or line.startswith('二、') or line.startswith('三、') or line.startswith('四、'):
            p = doc.add_paragraph(line)
            p.runs[0].bold = True
        elif line.startswith('（一）') or line.startswith('（二）') or line.startswith('（三）') or line.startswith('（四）'):
            p = doc.add_paragraph(line)
            p.runs[0].bold = True
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Pt(24)

    # 页脚
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "—— 「免费煎蛋卷（313版）」制作 ——"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # 保存
    home = os.path.expanduser('~')
    output_dir = os.path.join(home, '.claude', 'skills', 'essay-grader', 'outputs')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, '中世纪德国分裂原因批改报告.docx')
    doc.save(output_path)
    print(f"Word已生成: {output_path}")
    print("总分: 26.5/40 (三类卷)")

if __name__ == "__main__":
    main()
