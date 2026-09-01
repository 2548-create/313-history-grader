#!/usr/bin/env python3
"""
essay-grader: 生成313历史学论述题批改Word报告
完整流程：生成参考答案 → 提取踩分点 → OCR识别学生答案 → 批改 → 输出Word
"""
import os
import re
import sys
import json
import argparse
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ==================== 路径与配置 ====================
# skill 根目录：scripts/generate_word.py 的上级目录（不依赖运行时 CWD）
SKILL_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DEFAULT_OUTPUT_DIR = os.path.join(SKILL_ROOT, "outputs")
CONFIG_PATH = os.path.join(SKILL_ROOT, "config.json")
CONFIG_EXAMPLE_PATH = os.path.join(SKILL_ROOT, "config.example.json")


def _sanitize_filename(s, max_len=20):
    """清洗字符串为安全的文件名片段（去非法字符、空白转下划线、截断）"""
    s = re.sub(r'[\\/:*?"<>|\r\n\t]+', '', str(s)).strip()
    s = re.sub(r'\s+', '_', s)
    return s[:max_len]


def load_config():
    """读取配置：config.example.json 作默认模板，config.json 作用户覆盖（均不存在则回退内置默认）"""
    cfg = {"output_dir": "", "filename_template": "{topic}_{timestamp}.docx"}
    for p in (CONFIG_EXAMPLE_PATH, CONFIG_PATH):
        if os.path.exists(p):
            try:
                with open(p, 'r', encoding='utf-8') as f:
                    cfg.update(json.load(f))
            except Exception:
                pass
    return cfg


def resolve_output_path(input_json_path, config, cli_output=None):
    """解析最终输出路径。优先级：cli_output > config.output_dir > 默认(skill/outputs/)。
    - cli_output：相对 CWD（用户直觉）
    - config.output_dir：相对路径以 skill 根为基准解析
    返回绝对路径。"""
    if cli_output:
        return os.path.abspath(cli_output)
    out_dir = DEFAULT_OUTPUT_DIR
    cfg_dir = (config.get("output_dir") or "").strip()
    if cfg_dir:
        out_dir = cfg_dir if os.path.isabs(cfg_dir) else os.path.join(SKILL_ROOT, cfg_dir)
    template = config.get("filename_template") or "{topic}_{timestamp}.docx"
    topic = ""
    try:
        with open(input_json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        topic = _sanitize_filename(data.get("question", ""), 20)
    except Exception:
        pass
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    fname = template.format(topic=topic or "批改报告", timestamp=ts)
    if not fname.lower().endswith(".docx"):
        fname += ".docx"
    return os.path.join(out_dir, fname)


# ==================== AI腔检测 ====================
AI_VOICE_PATTERNS = [
    r'受.*驱动', r'遵循.*的逻辑', r'核心矛盾在于', r'结构性变化',
    r'双重驱动', r'根本转变', r'历史必然性', r'底层逻辑', r'结构性弱点',
    r'规模效应', r'制度效率', r'增长模式', r'不可持续性',
    r'深刻地', r'极大地', r'显著地', r'前所未有地', r'至关重要',
    r'完全.*废除', r'彻底.*改变', r'永远.*',
]

def check_ai_voice(text):
    """检测AI腔禁用词"""
    issues = []
    for pattern in AI_VOICE_PATTERNS:
        matches = re.findall(pattern, text)
        if matches:
            issues.extend(matches)
    return issues


# ==================== 分数校验 ====================
def check_score_consistency(sections, language_score, points_score):
    """检查分数一致性"""
    errors = []
    for section in sections:
        section_name = section.get('name', '未知板块')
        max_score = section.get('max_score', 0)
        score = section.get('score', 0)
        points = section.get('points', [])

        points_total = sum(float(p[3]) for p in points)
        if points_total != max_score:
            errors.append(f"板块'{section_name}'的踩分点分值之和({points_total})不等于板块满分({max_score})")

        student_total = sum(float(p[5]) for p in points)
        if student_total != score:
            errors.append(f"板块'{section_name}'的学生得分之和({student_total})不等于板块得分({score})")

    lang_total = language_score
    if lang_total > 10.01:
        errors.append(f"论述组织+卷面得分 {lang_total} 超过满分10分")

    total = points_score + lang_total
    if total > 40.01:
        errors.append(f"总分{total}超过满分40分")

    return errors


# ==================== 点评格式校验 ====================
def check_comment_format(comment_text, point_name):
    """检查点评格式"""
    errors = []
    if '【' in comment_text or '】' in comment_text:
        errors.append(f"踩分点'{point_name}'的点评包含【】标签")
    return errors


def check_comment_elements(comment_text, point_name):
    """检查点评四要素"""
    warnings = []
    if not any(kw in comment_text for kw in ['答出', '命中', '完整', '基本', '部分', '触及']):
        warnings.append(f"踩分点'{point_name}'的点评可能缺少'答出了什么'要素")
    if not any(kw in comment_text for kw in ['遗漏', '错误', '缺失', '未提及', '扣']) and '—' not in comment_text:
        warnings.append(f"踩分点'{point_name}'的点评可能缺少'哪里有问题'要素")
    if not any(kw in comment_text for kw in ['扣', '分']) and '—' not in comment_text:
        warnings.append(f"踩分点'{point_name}'的点评可能缺少'扣分'要素")
    if not any(kw in comment_text for kw in ['正确表述', '应为']) and '—' not in comment_text:
        warnings.append(f"踩分点'{point_name}'的点评可能缺少'正确表述'要素")
    return warnings


# ==================== 学生作答格式校验 ====================
def check_student_answer_column(student_texts):
    """检查学生作答列格式"""
    errors = []
    for i, text in enumerate(student_texts):
        if text and ('学生答出' in text or '学生写了' in text or '学生答' in text):
            errors.append(f"学生作答列第{i+1}行包含引导语")
        if text and text.startswith('"') and text.endswith('"'):
            errors.append(f"学生作答列第{i+1}行使用双引号包裹")
    return errors


# ==================== 整体点评格式校验 ====================
def check_diagnosis_format(diagnosis_text):
    """检查整体点评格式"""
    errors = []
    warnings = []
    if diagnosis_text:
        if re.search(r'属.*类卷.*（.*）', diagnosis_text):
            errors.append("整体点评使用括号内嵌理由")
        elements = {
            '等级理由': r'属于?[一二三类]类卷',
            '得分': r'\d+\.?\d*/40',
            '要点踩分': r'要点|踩分|史实',
            '组织论述': r'结构|逻辑|论述|组织',
            '亮点': r'亮点|较好|完整',
            '不足': r'失分|不足|缺失|错误',
            '提分空间': r'若能在|有望|水平'
        }
        missing = [k for k, v in elements.items() if not re.search(v, diagnosis_text)]
        if missing:
            warnings.append(f"整体点评可能缺少要素：{', '.join(missing)}")
    return errors, warnings


# ==================== 参考答案格式校验 ====================
def check_reference_answer(ref_text):
    """检查参考答案格式"""
    errors = []
    warnings = []
    if ref_text:
        word_count = len(ref_text.replace('\n', '').replace(' ', ''))
        if word_count < 1000:
            errors.append(f"参考答案字数{word_count}不足1000字")
        ai_issues = check_ai_voice(ref_text)
        if ai_issues:
            errors.append(f"参考答案包含AI腔禁用词: {', '.join(set(ai_issues))}")
        if not ref_text.startswith('  '):
            errors.append("参考答案开篇应首行缩进两格（背景段）")
        # 检查英文残留
        if re.search(r'[a-zA-Z]{2,}', ref_text):
            errors.append("参考答案包含英文残留，请检查")
    return errors, warnings


class EssayGrader:
    def __init__(self):
        self.question = None
        self.student_answer = None
        self.red_keywords = []
        self.reference_answer = None
        self.scores = {'points': 0, 'language': 0, '卷面': 0}
        self.sections = []
        self.language_breakdown = {}
        self.issues = []
        self.suggestions = []
        self.diagnosis = None
        self.reference = None

    def set_question(self, question):
        self.question = question

    def set_student_answer(self, answer, red_keywords=None):
        self.student_answer = answer
        self.red_keywords = []
        if red_keywords:
            for kw in red_keywords:
                if isinstance(kw, tuple) and len(kw) == 2:
                    self.red_keywords.append(kw)

    def _add_answer_paragraph(self, doc, text):
        """添加学生答案段落，标红史实硬伤"""
        p = doc.add_paragraph()
        if not self.red_keywords:
            run = p.add_run(text)
        else:
            segments = [text]
            for error_word, correct_word in self.red_keywords:
                new_segments = []
                for seg in segments:
                    if isinstance(seg, tuple):
                        new_segments.append(seg)
                        continue
                    parts = seg.split(error_word)
                    for i, part in enumerate(parts):
                        if part:
                            new_segments.append(part)
                        if i < len(parts) - 1:
                            new_segments.append((error_word, correct_word))
                segments = new_segments
            for seg in segments:
                if isinstance(seg, tuple):
                    error_word, correct_word = seg
                    run = p.add_run(error_word)
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.bold = True
                    if correct_word:
                        note = p.add_run(f"（应为{correct_word}）")
                        note.font.color.rgb = RGBColor(255, 0, 0)
                        note.font.size = Pt(9)
                else:
                    run = p.add_run(seg)

    def set_grading_result(self, result):
        """设置完整批改结果"""
        self.sections = result.get('sections', [])
        self.language_breakdown = result.get('language_breakdown', {})
        self.issues = result.get('issues', [])
        self.suggestions = result.get('suggestions', [])
        self.reference = result.get('reference', '')
        self.diagnosis = result.get('diagnosis', '') or None

        # 从表格数据读取分数，确保一致性
        self.scores['points'] = sum(s.get('score', 0) for s in self.sections)
        self.scores['language'] = sum(v.get('score', 0) for v in self.language_breakdown.values())
        self.scores['卷面'] = result.get('卷面', self.language_breakdown.get('neatness', {}).get('score', 0))

        # 强制同步diagnosis分数
        total = self.scores['points'] + self.scores['language']
        import re
        if self.diagnosis:
            # 替换diagnosis中的分数为计算值
            self.diagnosis = re.sub(r'\d+\.?\d*/40', f'{total}/40', self.diagnosis)

        # 验证语言组织评分不超过满分
        lang_total = self.scores['language'] + self.scores['卷面']
        if lang_total > 10.01:
            ratio = 10.0 / lang_total
            self.scores['language'] = round(self.scores['language'] * ratio, 1)
            self.scores['卷面'] = round(self.scores['卷面'] * ratio, 1)
            # 重新同步diagnosis
            total = self.scores['points'] + self.scores['language']
            self.diagnosis = re.sub(r'\d+\.?\d*/40', f'{total}/40', self.diagnosis)

    def run_all_checks(self):
        """运行所有检查，返回(errors, warnings)"""
        all_errors = []
        all_warnings = []

        # 检查参考答案
        ref_errors, ref_warnings = check_reference_answer(self.reference)
        all_errors.extend(ref_errors)
        all_warnings.extend(ref_warnings)

        # 检查点评格式和四要素
        for section in self.sections:
            for point in section.get('points', []):
                comment = point[6] if len(point) > 6 else ''
                point_name = point[2] if len(point) > 2 else '未知'
                fmt_errors = check_comment_format(comment, point_name)
                elem_warnings = check_comment_elements(comment, point_name)
                all_errors.extend(fmt_errors)
                all_warnings.extend(elem_warnings)

        # 检查分数一致性
        score_errors = check_score_consistency(
            self.sections, self.scores['language'] + self.scores['卷面'], self.scores['points']
        )
        all_errors.extend(score_errors)

        # 检查学生作答列
        student_texts = []
        for section in self.sections:
            for point in section.get('points', []):
                if len(point) > 2:
                    student_texts.append(point[2])
        ans_errors = check_student_answer_column(student_texts)
        all_errors.extend(ans_errors)

        # 检查整体点评
        diag_errors, diag_warnings = check_diagnosis_format(self.diagnosis)
        all_errors.extend(diag_errors)
        all_warnings.extend(diag_warnings)

        return all_errors, all_warnings

    def generate_word_report(self, output_path):
        """生成Word版批改报告"""
        # 先运行所有检查
        errors, warnings = self.run_all_checks()
        if errors:
            print("错误：")
            for e in errors:
                print(f"  - {e}")
            print("请修正错误后重新生成。")
            return None
        if warnings:
            print("警告：")
            for w in warnings:
                print(f"  - {w}")

        doc = Document()
        doc.styles['Normal'].font.name = 'SimSun'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(10)

        # 标题
        title = doc.add_heading('313历史学统考论述题批改报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 题目
        doc.add_heading('【题目】', level=1)
        doc.add_paragraph(self.question)

        # 总分概览
        doc.add_heading('【总分概览】', level=1)
        total_score = self.scores['points'] + self.scores['language'] + self.scores['卷面']
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Light Grid Accent 1'
        for i, h in enumerate(['维度', '得分', '满分']):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
        for i, row in enumerate([
            ['要点踩点', str(self.scores['points']), '30'],
            ['论述组织（含卷面）', str(self.scores['language'] + self.scores['卷面']), '10'],
            ['合计', f'{total_score}/40', '']
        ]):
            for j, cell_text in enumerate(row):
                table.rows[i + 1].cells[j].text = cell_text
        doc.add_paragraph()

        # 学生答案
        doc.add_heading('【学生答案】', level=1)
        if isinstance(self.student_answer, str):
            for para_text in self.student_answer.split('\n\n'):
                para_text = para_text.strip()
                if para_text:
                    self._add_answer_paragraph(doc, para_text)
        else:
            self._add_answer_paragraph(doc, str(self.student_answer))
        if self.red_keywords:
            note_p = doc.add_paragraph()
            note_run = note_p.add_run("注：标红处为史实错误，括号内为正确表述。")
            note_run.font.color.rgb = RGBColor(128, 128, 128)
            note_run.font.size = Pt(9)
        doc.add_paragraph()

        # 踩分点对照表
        doc.add_heading('【踩分点对照表】', level=1)
        for section in self.sections:
            title_text = f"（{section['name']}）（满分{section['max_score']}分，得分{section['score']}分）"
            p = doc.add_paragraph(title_text)
            for run in p.runs:
                run.bold = True
            headers = ['踩分点', '学生作答', '得分/满分', '点评']
            n_rows = len(section['points']) + 1
            table = doc.add_table(rows=n_rows, cols=4)
            table.style = 'Light Grid Accent 1'
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
            for i, row in enumerate(section['points']):
                # p[0]=踩分点, p[2]=学生作答, p[3]=满分, p[5]=得分, p[6]=点评
                table.rows[i + 1].cells[0].text = str(row[0]) if len(row) > 0 else ''
                table.rows[i + 1].cells[1].text = str(row[2]) if len(row) > 2 else ''
                # 得分/满分列
                student_score = row[5] if len(row) > 5 else 0
                max_score = row[3] if len(row) > 3 else 0
                table.rows[i + 1].cells[2].text = f"{student_score}/{max_score}"
                table.rows[i + 1].cells[3].text = str(row[6]) if len(row) > 6 else ''
                for j in range(4):
                    for p in table.rows[i + 1].cells[j].paragraphs:
                        for run in p.runs:
                            run.bold = False
            doc.add_paragraph()

        # 论述组织评分
        doc.add_heading('【论述组织评分】（满分10分）', level=1)
        lang_dims = [
            ('整体结构', 'structure', 2),
            ('论证逻辑', 'logic', 2),
            ('史论结合', 'argument', 2),
            ('文字表达', 'language', 1),
            ('史观运用', 'history_view', 1),
            ('卷面整洁', 'neatness', 2),
        ]
        table_data = [['维度', '得分/满分', '评语']]
        total_lang = 0
        for name, key, max_score in lang_dims:
            info = self.language_breakdown.get(key, {})
            score = info.get('score', 0)
            reason = info.get('reason', info.get('suggestion', ''))
            table_data.append([name, f"{score}/{max_score}", reason])
            total_lang += score
        table_data.append(['合计', f"{total_lang}/10", ''])
        n_rows = len(table_data)
        table = doc.add_table(rows=n_rows, cols=3)
        table.style = 'Light Grid Accent 1'
        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = str(cell_text)
                if i == 0 or i == n_rows - 1:
                    cell.paragraphs[0].runs[0].bold = True
                else:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = False
        doc.add_paragraph()

        # 本题诊断
        doc.add_heading('【本题诊断】', level=1)
        doc.add_paragraph('一、亮点', style='Heading 2')
        highlights = [i for i in self.issues if i.get('type') == '亮点']
        for issue in highlights:
            doc.add_paragraph(f"• {issue.get('problem', '')}")
        if not highlights:
            doc.add_paragraph("• 无明显亮点")
        doc.add_paragraph('二、不足', style='Heading 2')
        for issue in self.issues:
            if issue.get('type') != '亮点':
                p = doc.add_paragraph(f"• {issue.get('problem', '')}")
                fix = issue.get('fix', '')
                if fix:
                    p.add_run(f"  正确表述：{fix}")
        if not any(i.get('type') != '亮点' for i in self.issues):
            doc.add_paragraph("• 无明显不足")
        doc.add_paragraph('三、改进方向', style='Heading 2')
        for s in self.suggestions:
            if isinstance(s, dict):
                doc.add_paragraph(f"• {s.get('problem', '')}")
            else:
                doc.add_paragraph(f"• {s}")
        if not self.suggestions:
            doc.add_paragraph("• 暂无改进建议")
        doc.add_paragraph()

        # 整体点评
        doc.add_heading('【整体点评】', level=1)
        if self.diagnosis:
            doc.add_paragraph(self.diagnosis)

        # 参考答案
        doc.add_heading('【参考答案】', level=1)
        if self.reference:
            for line in self.reference.split('\n'):
                line = line.strip()
                if not line:
                    continue
                if re.match(r'^[一二三四五六七八九十]+、', line) or re.match(r'^（[一二三四五六七八九十]+）', line):
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.bold = True
                else:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.first_line_indent = Pt(24)

        # 页脚
        footer = doc.sections[0].footer.paragraphs[0]
        footer.text = "—— 「免费煎蛋卷（313版）」制作 ——"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)

        d = os.path.dirname(output_path)
        if d:
            os.makedirs(d, exist_ok=True)
        doc.save(output_path)
        print(f"Word已生成: {output_path}")
        return output_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="生成313历史学论述题批改Word报告")
    parser.add_argument("input", help="批改数据 JSON 路径")
    parser.add_argument("-o", "--output", help="输出 docx 路径（绝对或相对CWD，优先级最高）")
    parser.add_argument("--config", help="自定义 config.json 路径（默认读 skill 根/config.json）")
    args = parser.parse_args()

    config = load_config()
    if args.config:
        cfg_path = os.path.abspath(os.path.expanduser(args.config))
        if os.path.exists(cfg_path):
            try:
                with open(cfg_path, 'r', encoding='utf-8') as f:
                    config.update(json.load(f))
            except Exception:
                pass

    output_path = resolve_output_path(args.input, config, args.output)
    with open(args.input, 'r', encoding='utf-8') as f:
        data = json.load(f)
    gen = EssayGrader()
    gen.set_grading_result(data)
    gen.generate_word_report(output_path)
